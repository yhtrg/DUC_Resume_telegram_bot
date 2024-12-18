import requests
import json
import os
import textract
import logging
import asyncio
from aiogram import Bot, Dispatcher, types
from aiogram.filters import CommandStart
from docx import Document as DocxDocument
from aiogram.types.input_file import FSInputFile
from docx.shared import Pt
from striprtf.striprtf import rtf_to_text
from pypdf import PdfReader
from spire.doc import *
from spire.doc.common import *
from docxtpl import DocxTemplate


logging.basicConfig(level=logging.INFO)

encoding = 'UTF-8'
API_URL = "http://178.212.132.7:3003/api/v1/prediction/914536a7-cb64-4504-afab-4d09c557e524"
API_URL_SPACES = "http://178.212.132.7:3003/api/v1/prediction/3a8b3a53-84b6-4551-8127-7162ca12cb64"
TOKEN = '7096081921:AAHX23mpdT1pe4yZJfzBxnNM10xroSkB8HI'
#TOKEN = '7266662655:AAG5CUk3OSiylrlM6QfD_XiL4yeGB7o8AjQ'
DOWNLOAD_FOLDER = 'downloads'
os.makedirs(DOWNLOAD_FOLDER, exist_ok=True)

dp = Dispatcher()
bot = Bot(token=TOKEN)

def query(payload):
    response = requests.post(API_URL, json=payload)
    return response.json()

def query_spaces(payload):
    response = requests.post(API_URL_SPACES, json=payload)
    return response.json()

def find_spaces(text):
    space_count = text.count(' ')
    char_count = len(text) - space_count
    
    return char_count, space_count

def download_word(data):
    tpl = DocxTemplate('resume.docx')
    
    for key, value in data.items():
        if key == 'name':
            file_name = f"{value.replace(' ', '_')}_Резюме.docx"
            break
        else:
            file_name = 'Резюме.docx'
                
    tpl.render(data)
    tpl.save(os.path.join(DOWNLOAD_FOLDER, file_name))
    return file_name

@dp.message(CommandStart())
async def start_command(message: types.Message):
    await message.reply("Привет! Отправь мне документ, и я создам резюме с ним.")

def process_document(file_path):
    response = ''
    if '.pdf' in file_path:
        reader = PdfReader(file_path)
        for page in reader.pages:
            response += page.extract_text() + '\n'

    elif '.docx' in file_path:
        doc = DocxDocument(file_path)
        for paragraph in doc.paragraphs:
            response += paragraph.text + '\n'

        if doc.tables:
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    response += ' '.join(row_text)

    elif '.rtf' in file_path:
        with open(file_path) as infile:
            content = infile.read()
            response = rtf_to_text(content)

    elif '.doc' in file_path:
        output_file = file_path[:-5] + '.docx'
        document = Document()
        document.LoadFromFile(file_path)
        document.SaveToFile(output_file, FileFormat.Docx2016)
        response = textract.process(output_file).decode(encoding)
        os.remove(output_file)
    return response

@dp.message()
async def handle_document(message: types.Message):
    user_id = message.from_user.id
    input_document = message.document
    input_file_name = input_document.file_name
    input_file_path = os.path.join(DOWNLOAD_FOLDER, input_file_name)
    input_file_id = input_document.file_id
    input_file = await bot.get_file(input_file_id)

    await bot.download_file(input_file.file_path, input_file_path)
    await message.reply("Документ получен! Обрабатываю файл...")

    response = process_document(input_file_path)
    cleaned_response = repr(response.replace('\n', ' ').replace('\r', ''))
    print(cleaned_response, type(cleaned_response))
    char_count, space_count = find_spaces(cleaned_response)
    if (space_count * 100)/(char_count + space_count) >= 20:
        cleaned_response.replace(' ', '')
        cleaned_response = query_spaces({"question": cleaned_response})
        output_flowise = query({"question": cleaned_response['text']})
    else:
        output_flowise = query({"question" : cleaned_response})
    
    if '```json' in output_flowise['text']:
        try:
            data = json.loads(output_flowise['text'][7:-3])
        except:
            data = json.loads(output_flowise['text'][7:-4])
    else:
        data = json.loads(output_flowise['text'])
    print(data)

    output_file_name = download_word(data)
    output_document = FSInputFile(f'downloads/{output_file_name}')
    await bot.send_document(user_id, output_document)
    
    os.remove(os.path.join(DOWNLOAD_FOLDER, input_file_name))
    os.remove(os.path.join(DOWNLOAD_FOLDER, output_file_name))

async def main() -> None:
    await dp.start_polling(bot)

if __name__ == '__main__':
    logging.basicConfig(level=logging.INFO)
    asyncio.run(main())
